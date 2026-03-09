"""
工程师智能体 — 后端路由
第一个子能力：Doc2Audio（文档转语音）
"""
import os
import uuid
import re
import asyncio
from pathlib import Path

from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
from docx import Document

router = APIRouter(prefix="/api/engineer", tags=["engineer"])

# ---- 目录配置 ----
_BASE_DIR = Path(__file__).resolve().parents[2]  # memoryhub/backend
UPLOAD_DIR = _BASE_DIR / "data" / "engineer_uploads"
OUTPUT_DIR = _BASE_DIR / "data" / "engineer_output"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {"doc", "docx"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


# =====================================================================
# 工具函数（从 doc2audio/app.py 迁移）
# =====================================================================

def _allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def _extract_text_from_docx(file_path: str) -> str:
    """从 DOCX 文件中提取文本"""
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)
    return "\n".join(text)


def _extract_text_from_doc(file_path: str) -> str:
    """从 DOC 文件中提取文本（优先 antiword，回退二进制提取）"""
    try:
        import subprocess
        result = subprocess.run(["antiword", file_path], capture_output=True, text=True)
        if result.returncode == 0:
            return result.stdout
    except FileNotFoundError:
        pass

    # 回退：提取可见字符
    with open(file_path, "rb") as f:
        content = f.read()
        text = ""
        i = 0
        while i < len(content):
            if content[i] >= 0x80:
                try:
                    char = content[i : i + 3].decode("utf-8")
                    if char.strip():
                        text += char
                    i += 3
                    continue
                except Exception:
                    pass
            elif 32 <= content[i] <= 126 or content[i] in (9, 10, 13):
                text += chr(content[i])
            i += 1
        return text


def _detect_language(text: str) -> str:
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    english_chars = len(re.findall(r"[a-zA-Z]", text))
    total = chinese_chars + english_chars
    if total == 0:
        return "en"
    return "zh" if chinese_chars / total > 0.3 else "en"


async def _synthesize_speech(text: str, output_file: str, lang: str = "zh-CN", speed: str = "normal", volume: str = "normal"):
    """使用 edge-tts 合成语音（原生 async）"""
    import edge_tts

    voice_map = {
        "zh-CN": "zh-CN-XiaoxiaoNeural",
        "en": "en-US-JennyNeural",
    }
    voice = voice_map.get(lang, "zh-CN-XiaoxiaoNeural")

    rate_map = {"slow": "-20%", "normal": "+0%", "fast": "+20%"}
    rate = rate_map.get(speed, "+0%")

    volume_map = {"low": "-30%", "normal": "+0%", "loud": "+30%"}
    vol = volume_map.get(volume, "+0%")

    communicate = edge_tts.Communicate(text, voice, rate=rate, volume=vol)
    await communicate.save(output_file)


# =====================================================================
# API 路由
# =====================================================================

@router.get("/doc2audio/voices")
async def get_voices():
    """返回可用语音列表"""
    return {
        "zh": [
            {"id": "xiaoxiao", "name": "晓晓", "description": "中文女声（推荐）"},
            {"id": "xiaoyi", "name": "晓伊", "description": "中文女声"},
            {"id": "yunyang", "name": "云扬", "description": "中文男声"},
        ],
        "en": [
            {"id": "jenny", "name": "Jenny", "description": "英文女声（推荐）"},
            {"id": "aria", "name": "Aria", "description": "英文女声"},
            {"id": "guy", "name": "Guy", "description": "英文男声"},
        ],
    }


@router.post("/doc2audio")
async def convert_doc_to_audio(
    file: UploadFile = File(...),
    voice: str = Form("auto"),
    speed: str = Form("normal"),
    volume: str = Form("normal"),
):
    """上传 DOC/DOCX 文件 → 转换为 MP3"""
    if not file.filename or not _allowed_file(file.filename):
        raise HTTPException(status_code=400, detail="不支持的文件格式，请上传 DOC 或 DOCX 文件")

    # 读取并保存上传文件
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="文件大小不能超过 50MB")

    file_ext = file.filename.rsplit(".", 1)[1].lower()
    unique_id = str(uuid.uuid4())
    upload_path = UPLOAD_DIR / f"{unique_id}.{file_ext}"
    upload_path.write_bytes(content)

    try:
        # 提取文本
        if file_ext == "docx":
            text = _extract_text_from_docx(str(upload_path))
        else:
            text = _extract_text_from_doc(str(upload_path))

        if not text or not text.strip():
            raise HTTPException(status_code=422, detail="文件中没有找到文本内容")

        # 检测语言
        language = _detect_language(text)
        lang = "zh-CN" if language == "zh" else "en"

        # 合成语音
        output_filename = f"{unique_id}.mp3"
        output_path = OUTPUT_DIR / output_filename
        await _synthesize_speech(text, str(output_path), lang, speed, volume)

        return {
            "success": True,
            "download_url": f"/api/engineer/doc2audio/download/{output_filename}",
            "language": language,
            "voice": lang,
            "text_length": len(text),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")
    finally:
        # 清理上传文件
        try:
            upload_path.unlink(missing_ok=True)
        except Exception:
            pass


@router.get("/doc2audio/download/{filename}")
async def download_audio(filename: str):
    """下载生成的 MP3 文件"""
    file_path = OUTPUT_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(
        path=str(file_path),
        media_type="audio/mpeg",
        filename=f"audio_{filename}",
    )
